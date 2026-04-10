"""
文档解析器 - 支持多种格式的文档解析
"""

import os
import logging
from pathlib import Path
from typing import Optional, List, Dict, Any
import tempfile

logger = logging.getLogger(__name__)


class DocumentParser:
    """文档解析器，支持多种格式"""

    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def parse_document(self, file_path: str) -> str:
        """
        解析文档文件，提取文本内容

        Args:
            file_path: 文档文件路径

        Returns:
            提取的文本内容
        """
        try:
            file_ext = Path(file_path).suffix.lower()

            if file_ext == '.pdf':
                return self._parse_pdf(file_path)
            elif file_ext in ['.doc', '.docx']:
                return self._parse_word(file_path)
            elif file_ext in ['.xls', '.xlsx']:
                return self._parse_excel(file_path)
            elif file_ext in ['.txt', '.md', '.json', '.yaml', '.yml']:
                return self._parse_text_file(file_path)
            elif file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff']:
                return self._parse_image(file_path)
            else:
                # 尝试作为文本文件读取
                return self._parse_text_file(file_path)

        except Exception as e:
            self.logger.error(f"解析文档失败 {file_path}: {e}")
            # 返回错误信息
            return f"[文档解析失败: {str(e)}]"

    def _parse_pdf(self, file_path: str) -> str:
        """解析PDF文件"""
        try:
            # 优先使用PyMuPDF（性能最好）
            try:
                import fitz  # PyMuPDF
                text = ""
                doc = fitz.open(file_path)
                for page_num in range(doc.page_count):
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    if page_text:
                        text += page_text + "\n\n"
                doc.close()
                if text.strip():
                    return text.strip()
            except ImportError:
                self.logger.debug("PyMuPDF未安装，尝试其他方法")
            except Exception as e:
                self.logger.warning(f"PyMuPDF解析失败: {e}")

            # 其次使用pdfplumber（表格提取更好）
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"

                        # 尝试提取表格
                        tables = page.extract_tables()
                        if tables:
                            for table_num, table in enumerate(tables):
                                if table:
                                    text += f"\n[表格 {table_num + 1}]\n"
                                    for row in table:
                                        if any(cell is not None for cell in row):
                                            text += " | ".join(str(cell) if cell is not None else "" for cell in row) + "\n"
                                    text += "\n"
                return text.strip()
            except ImportError:
                self.logger.debug("pdfplumber未安装，尝试其他方法")
            except Exception as e:
                self.logger.warning(f"pdfplumber解析失败: {e}")

            # 最后使用PyPDF2
            try:
                import PyPDF2
                text = ""
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
                return text.strip()
            except ImportError:
                self.logger.debug("PyPDF2未安装")
            except Exception as e:
                self.logger.warning(f"PyPDF2解析失败: {e}")

            # 如果所有库都未安装，返回提示信息
            return f"[PDF解析失败: 请安装PyMuPDF、pdfplumber或PyPDF2库]\n文件路径: {file_path}\n建议安装: pip install pymupdf pdfplumber PyPDF2"

        except Exception as e:
            self.logger.error(f"PDF解析失败: {e}")
            return f"[PDF解析失败: {str(e)}]"

    def _parse_word(self, file_path: str) -> str:
        """解析Word文档"""
        try:
            file_ext = Path(file_path).suffix.lower()

            # 对于.docx文件，使用python-docx
            if file_ext == '.docx':
                try:
                    import docx
                    doc = docx.Document(file_path)
                    text = ""

                    # 提取段落
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"

                    # 提取表格
                    for table in doc.tables:
                        text += "\n[表格]\n"
                        for row in table.rows:
                            row_text = " | ".join(cell.text.strip() for cell in row.cells)
                            if row_text.strip():
                                text += row_text + "\n"
                        text += "\n"

                    return text.strip()
                except ImportError:
                    self.logger.warning("python-docx未安装")
                except Exception as e:
                    self.logger.warning(f"python-docx解析失败: {e}")

            # 对于.doc文件，尝试多种方法
            elif file_ext == '.doc':
                # 尝试使用antiword（需要系统安装）
                try:
                    import subprocess
                    result = subprocess.run(
                        ['antiword', file_path],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    if result.returncode == 0:
                        return result.stdout.strip()
                except (FileNotFoundError, subprocess.TimeoutExpired):
                    self.logger.debug("antiword未安装或执行失败")

                # 尝试使用catdoc
                try:
                    import subprocess
                    result = subprocess.run(
                        ['catdoc', file_path],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    if result.returncode == 0:
                        return result.stdout.strip()
                except (FileNotFoundError, subprocess.TimeoutExpired):
                    self.logger.debug("catdoc未安装或执行失败")

            # 通用提示
            return f"[Word文档解析失败: 请安装python-docx库]\n文件路径: {file_path}\n对于.docx文件: pip install python-docx\n对于.doc文件: 需要安装antiword或catdoc工具"

        except Exception as e:
            self.logger.error(f"Word解析失败: {e}")
            return f"[Word解析失败: {str(e)}]"

    def _parse_excel(self, file_path: str) -> str:
        """解析Excel文件（提取全部文本内容，适合作为设计文档使用）"""
        try:
            import pandas as pd

            text_parts = []

            try:
                excel_file = pd.ExcelFile(file_path)
                sheet_names = excel_file.sheet_names

                for sheet_name in sheet_names:
                    try:
                        # 读取全部行
                        df = pd.read_excel(file_path, sheet_name=sheet_name)

                        sheet_text = f"=== Sheet: {sheet_name} ===\n"
                        sheet_text += f"形状: {df.shape[0]}行 x {df.shape[1]}列\n"

                        if not df.empty:
                            columns = [str(c) for c in df.columns.tolist()]

                            # Markdown 表格格式输出全部数据
                            sheet_text += "| " + " | ".join(columns) + " |\n"
                            sheet_text += "| " + " | ".join("---" for _ in columns) + " |\n"
                            for i in range(len(df)):
                                row_vals = []
                                for val in df.iloc[i].tolist():
                                    cell = str(val) if pd.notna(val) else ""
                                    # 管道符转义，防止破坏表格
                                    cell = cell.replace("|", "\\|")
                                    row_vals.append(cell)
                                sheet_text += "| " + " | ".join(row_vals) + " |\n"

                        text_parts.append(sheet_text)

                    except Exception as e:
                        text_parts.append(f"=== Sheet: {sheet_name} (读取失败: {str(e)}) ===\n")

            except Exception as e:
                text_parts.append(f"[Excel文件读取失败: {str(e)}]")

            return "\n".join(text_parts).strip()

        except ImportError:
            return f"[Excel解析失败: 请安装pandas库]\n文件路径: {file_path}"
        except Exception as e:
            self.logger.error(f"Excel解析失败: {e}")
            return f"[Excel解析失败: {str(e)}]"

    def _parse_text_file(self, file_path: str) -> str:
        """解析文本文件"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'latin-1', 'cp1252']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
                except Exception:
                    continue

            # 如果所有编码都失败，尝试二进制读取
            try:
                with open(file_path, 'rb') as f:
                    content = f.read()
                    return content.decode('utf-8', errors='ignore')
            except Exception as e:
                raise Exception(f"无法读取文件: {str(e)}")

        except Exception as e:
            self.logger.error(f"文本文件解析失败: {e}")
            return f"[文本文件解析失败: {str(e)}]"

    def _parse_image(self, file_path: str) -> str:
        """解析图像文件（OCR）"""
        try:
            # 尝试使用pytesseract
            try:
                import pytesseract
                from PIL import Image

                # 打开图像
                image = Image.open(file_path)

                # 使用OCR提取文本
                text = pytesseract.image_to_string(image, lang='chi_sim+eng')  # 中文+英文

                if text.strip():
                    return text.strip()
                else:
                    return f"[图像OCR未提取到文本]\n文件路径: {file_path}"

            except ImportError as e:
                self.logger.warning(f"OCR库未安装: {e}")
            except Exception as e:
                self.logger.warning(f"OCR处理失败: {e}")

            # 尝试使用easyocr
            try:
                import easyocr

                # 创建reader（中文+英文）
                reader = easyocr.Reader(['ch_sim', 'en'])

                # 执行OCR
                result = reader.readtext(file_path)

                # 提取文本
                text = "\n".join([item[1] for item in result])

                if text.strip():
                    return text.strip()

            except ImportError as e:
                self.logger.warning(f"EasyOCR未安装: {e}")
            except Exception as e:
                self.logger.warning(f"EasyOCR处理失败: {e}")

            return f"[图像OCR失败: 请安装OCR库]\n文件路径: {file_path}\n建议安装:\n1. Tesseract OCR: pip install pytesseract pillow\n   还需要安装Tesseract软件: https://github.com/UB-Mannheim/tesseract/wiki\n2. EasyOCR: pip install easyocr"

        except Exception as e:
            self.logger.error(f"图像解析失败: {e}")
            return f"[图像解析失败: {str(e)}]"

    def get_supported_formats(self) -> List[str]:
        """获取支持的文档格式"""
        return [
            '.pdf', '.doc', '.docx',  # 文档
            '.xls', '.xlsx', '.csv',  # 表格
            '.txt', '.md', '.json', '.yaml', '.yml',  # 文本
            '.html', '.htm',  # 网页
            '.png', '.jpg', '.jpeg', '.bmp', '.tiff',  # 图像（需要OCR）
        ]

    def is_supported_format(self, file_path: str) -> bool:
        """检查文件格式是否支持"""
        file_ext = Path(file_path).suffix.lower()
        return file_ext in self.get_supported_formats()


# 单例实例
_document_parser = None

def get_document_parser() -> DocumentParser:
    """获取文档解析器实例"""
    global _document_parser
    if _document_parser is None:
        _document_parser = DocumentParser()
    return _document_parser