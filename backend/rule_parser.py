"""
规则文件解析器 - 支持PDF、Word、Excel格式的规则文件解析
"""

import os
import re
import json
from typing import Dict, List, Any, Optional, Union
from pathlib import Path
from dataclasses import dataclass, field
import pandas as pd
import openpyxl

# 尝试导入PDF和Word处理库
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("警告: PyPDF2未安装，PDF解析功能将不可用")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx未安装，Word解析功能将不可用")

# 导入现有的Excel解析器
from excel_parser import IntelligentExcelParser


@dataclass
class ColumnRule:
    """列规则定义"""
    column_name: str
    data_source: str  # 数据来源，如"file1.xlsx!Sheet1!A列"
    calculation_rule: Optional[str] = None  # 计算规则，如"=SUM(A1:A10)"
    validation_rule: Optional[str] = None  # 验证规则，如">0"
    description: Optional[str] = None  # 列描述


@dataclass
class SheetRule:
    """工作表规则定义"""
    sheet_name: str
    columns: List[ColumnRule] = field(default_factory=list)
    data_range: Optional[str] = None  # 数据范围，如"A1:Z100"
    header_rows: Optional[int] = None  # 表头行数


@dataclass
class FileRule:
    """文件规则定义"""
    file_name: str
    sheets: List[SheetRule] = field(default_factory=list)
    description: Optional[str] = None


@dataclass
class RuleSet:
    """规则集合"""
    expected_file: FileRule  # 预期输出文件规则
    source_files: List[FileRule] = field(default_factory=list)  # 源文件规则
    mapping_rules: Dict[str, str] = field(default_factory=dict)  # 映射规则
    calculation_rules: Dict[str, str] = field(default_factory=dict)  # 计算规则


class RuleParser:
    """规则文件解析器基类"""

    def __init__(self):
        self.excel_parser = IntelligentExcelParser()

    def parse(self, file_path: str) -> RuleSet:
        """解析规则文件"""
        raise NotImplementedError("子类必须实现此方法")

    def _extract_column_rules(self, text: str) -> List[ColumnRule]:
        """从文本中提取列规则"""
        column_rules = []

        # 模式1: 列名: 数据来源 [计算规则]
        pattern1 = r'([^:\n]+):\s*([^[\n]+)(?:\[([^\]]+)\])?'
        matches1 = re.findall(pattern1, text)

        for match in matches1:
            column_name = match[0].strip()
            data_source = match[1].strip()
            calculation_rule = match[2].strip() if match[2] else None

            column_rules.append(ColumnRule(
                column_name=column_name,
                data_source=data_source,
                calculation_rule=calculation_rule
            ))

        # 模式2: 表格格式
        table_pattern = r'\|([^|]+)\|([^|]+)\|([^|]+)?\|'
        matches2 = re.findall(table_pattern, text)

        for match in matches2:
            if len(match) >= 2:
                column_name = match[0].strip()
                data_source = match[1].strip()
                calculation_rule = match[2].strip() if len(match) > 2 and match[2].strip() else None

                # 避免重复添加
                if not any(cr.column_name == column_name for cr in column_rules):
                    column_rules.append(ColumnRule(
                        column_name=column_name,
                        data_source=data_source,
                        calculation_rule=calculation_rule
                    ))

        return column_rules

    def _extract_sheet_rules(self, text: str) -> List[SheetRule]:
        """从文本中提取工作表规则"""
        sheet_rules = []

        # 查找工作表名称
        sheet_patterns = [
            r'工作表[：:]\s*([^\n]+)',
            r'Sheet[：:]\s*([^\n]+)',
            r'表[：:]\s*([^\n]+)',
            r'^([^:\n]+)[：:]\s*$'
        ]

        for pattern in sheet_patterns:
            sheet_matches = re.finditer(pattern, text, re.MULTILINE)
            for match in sheet_matches:
                sheet_name = match.group(1).strip()

                # 提取该工作表后面的内容
                start_pos = match.end()
                end_pos = text.find('\n\n', start_pos)
                if end_pos == -1:
                    end_pos = len(text)

                sheet_text = text[start_pos:end_pos]
                column_rules = self._extract_column_rules(sheet_text)

                if column_rules:
                    sheet_rules.append(SheetRule(
                        sheet_name=sheet_name,
                        columns=column_rules
                    ))

        return sheet_rules

    def _extract_file_rules(self, text: str) -> List[FileRule]:
        """从文本中提取文件规则"""
        file_rules = []

        # 查找文件名称
        file_patterns = [
            r'文件[：:]\s*([^\n]+)',
            r'File[：:]\s*([^\n]+)',
            r'^([^:\n]+)\.(xlsx?|xls|pdf|docx?)[：:]\s*$'
        ]

        for pattern in file_patterns:
            file_matches = re.finditer(pattern, text, re.MULTILINE)
            for match in file_matches:
                file_name = match.group(1).strip()

                # 提取该文件后面的内容
                start_pos = match.end()
                end_pos = text.find('\n\n', start_pos)
                if end_pos == -1:
                    end_pos = len(text)

                file_text = text[start_pos:end_pos]
                sheet_rules = self._extract_sheet_rules(file_text)

                if sheet_rules:
                    file_rules.append(FileRule(
                        file_name=file_name,
                        sheets=sheet_rules
                    ))

        return file_rules


class PDFRuleParser(RuleParser):
    """PDF规则文件解析器"""

    def __init__(self):
        super().__init__()
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2未安装，无法使用PDF解析功能")

    def parse(self, file_path: str) -> RuleSet:
        """解析PDF规则文件"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"PDF文件不存在: {file_path}")

        # 提取PDF文本
        text = self._extract_pdf_text(file_path)

        # 解析规则
        return self._parse_rules_from_text(text, file_path)

    def _extract_pdf_text(self, file_path: str) -> str:
        """从PDF文件中提取文本"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2未安装，无法解析PDF文件")

        text_parts = []

        try:
            with open(file_path, 'rb') as file:
                # 检查文件大小
                file.seek(0, 2)  # 移动到文件末尾
                file_size = file.tell()
                file.seek(0)  # 回到文件开头

                if file_size == 0:
                    raise ValueError(f"PDF文件为空: {file_path}")

                # 尝试读取PDF
                pdf_reader = PyPDF2.PdfReader(file)

                # 检查PDF是否加密
                if pdf_reader.is_encrypted:
                    # 尝试使用空密码解密
                    try:
                        pdf_reader.decrypt('')
                    except:
                        raise ValueError(f"PDF文件已加密，无法读取: {file_path}")

                total_pages = len(pdf_reader.pages)
                if total_pages == 0:
                    raise ValueError(f"PDF文件没有页面: {file_path}")

                for page_num in range(total_pages):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()

                        if not page_text or page_text.strip() == '':
                            # 可能是扫描的PDF，没有可提取的文本
                            text_parts.append(f"[第{page_num+1}页: 无文本内容，可能是扫描图像]")
                        else:
                            # 清理文本
                            cleaned_text = self._clean_text(page_text)
                            text_parts.append(cleaned_text)

                    except Exception as page_error:
                        text_parts.append(f"[第{page_num+1}页提取失败: {str(page_error)}]")

        except PyPDF2.errors.PdfReadError as e:
            raise ValueError(f"PDF文件格式错误: {str(e)}")
        except Exception as e:
            raise ValueError(f"读取PDF文件失败: {str(e)}")

        if not text_parts:
            return "[PDF文件没有可提取的文本内容]"

        result = '\n'.join(text_parts)

        # 检查提取的文本是否有效
        if len(result.strip()) < 10:  # 如果文本太短，可能提取失败
            return "[PDF文件可能包含扫描图像，无法提取文本]"

        return result

    def _clean_text(self, text: str) -> str:
        """清理提取的文本"""
        if not text:
            return ""

        # 1. 移除PDF中常见的特殊控制字符 (如\x01, \x02, \x03等)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', ' ', text)

        # 2. 替换PDF中常见的表格分隔符和特殊字符
        text = text.replace('', ' ')  # 替换SOH字符
        text = text.replace('', ' ')  # 替换STX字符
        text = text.replace('', ' ')  # 替换ETX字符

        # 3. 规范化空格和换行
        text = re.sub(r'\s+', ' ', text)

        # 4. 修复常见的中文标点问题
        text = text.replace(' ,', ',').replace(' .', '.').replace(' :', ':')
        text = text.replace(' ;', ';').replace(' ?', '?').replace(' !', '!')

        # 5. 修复中文引号
        text = text.replace(' "', '"').replace('" ', '"')
        text = text.replace(" '", "'").replace("' ", "'")

        # 6. 移除行首行尾的空格
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if line:  # 只保留非空行
                cleaned_lines.append(line)

        return '\n'.join(cleaned_lines)

    def _parse_rules_from_text(self, text: str, source_file: str) -> RuleSet:
        """从文本中解析规则"""
        # 首先检查是否是需求文档格式（如达美乐文档）
        if self._is_requirement_document(text):
            return self._parse_requirement_document(text, source_file)

        # 查找预期输出文件
        expected_file_patterns = [
            r'预期输出[：:]\s*([^\n]+)',
            r'Expected Output[：:]\s*([^\n]+)',
            r'Output File[：:]\s*([^\n]+)'
        ]

        expected_file_name = None
        for pattern in expected_file_patterns:
            match = re.search(pattern, text)
            if match:
                expected_file_name = match.group(1).strip()
                break

        if not expected_file_name:
            # 如果没有明确指定，使用源文件名作为基础
            base_name = Path(source_file).stem
            expected_file_name = f"{base_name}_output.xlsx"

        # 提取文件规则
        file_rules = self._extract_file_rules(text)

        # 分离预期文件和源文件
        expected_file_rule = None
        source_file_rules = []

        for file_rule in file_rules:
            if file_rule.file_name == expected_file_name or file_rule.file_name.endswith('_output.xlsx'):
                expected_file_rule = file_rule
            else:
                source_file_rules.append(file_rule)

        # 如果没有找到预期文件规则，创建一个
        if not expected_file_rule:
            expected_file_rule = FileRule(
                file_name=expected_file_name,
                sheets=[]
            )

        # 提取映射规则和计算规则
        mapping_rules = self._extract_mapping_rules(text)
        calculation_rules = self._extract_calculation_rules(text)

        return RuleSet(
            expected_file=expected_file_rule,
            source_files=source_file_rules,
            mapping_rules=mapping_rules,
            calculation_rules=calculation_rules
        )

    def _is_requirement_document(self, text: str) -> bool:
        """检查文本是否是需求文档格式"""
        # 检查是否包含需求文档常见关键词
        requirement_keywords = [
            '需求文档', '需求说明', '需求背景', '版本信息', '变更日志',
            '文档说明', '名词解释', '前期准备', '导入表说明', '报表项',
            '达美乐', '全职需求'
        ]

        text_lower = text.lower()
        for keyword in requirement_keywords:
            if keyword.lower() in text_lower:
                return True

        # 检查是否包含表格结构
        table_indicators = ['列名', '是否需要存储', '与报表项关系', '备注']
        for indicator in table_indicators:
            if indicator in text:
                return True

        return False

    def _parse_requirement_document(self, text: str, source_file: str) -> RuleSet:
        """解析需求文档格式的规则"""
        print(f"检测到需求文档格式，开始解析...")

        # 提取表格信息
        tables = self._extract_tables_from_text(text)

        # 创建文件规则
        source_file_rules = []
        expected_file_rule = None

        # 从表格中提取文件规则
        for table_name, table_data in tables.items():
            if '报表项' in table_name or '输出' in table_name:
                # 这是预期输出文件
                expected_file_rule = self._create_file_rule_from_table(table_name, table_data, is_expected=True)
            else:
                # 这是源文件
                file_rule = self._create_file_rule_from_table(table_name, table_data, is_expected=False)
                source_file_rules.append(file_rule)

        # 如果没有找到预期文件，创建一个默认的
        if not expected_file_rule:
            base_name = Path(source_file).stem
            expected_file_name = f"{base_name}_output.xlsx"
            expected_file_rule = FileRule(
                file_name=expected_file_name,
                sheets=[]
            )

        # 提取映射规则（从表格中的"与报表项关系"列）
        mapping_rules = self._extract_mapping_from_tables(tables)

        # 提取计算规则（从文本中的公式）
        calculation_rules = self._extract_calculation_rules(text)

        print(f"需求文档解析完成:")
        print(f"  找到 {len(source_file_rules)} 个源文件规则")
        print(f"  找到 {len(mapping_rules)} 个映射规则")
        print(f"  找到 {len(calculation_rules)} 个计算规则")

        return RuleSet(
            expected_file=expected_file_rule,
            source_files=source_file_rules,
            mapping_rules=mapping_rules,
            calculation_rules=calculation_rules
        )

    def _extract_tables_from_text(self, text: str) -> Dict[str, List[Dict[str, str]]]:
        """从文本中提取表格数据"""
        tables = {}

        # 分割文本为行
        lines = text.split('\n')

        current_table = None
        current_headers = []
        current_rows = []

        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # 检查是否是表头行（包含"列名"、"是否需要存储"等）
            if '列名' in line and ('是否需要存储' in line or '与报表项关系' in line):
                # 找到表格开始
                if current_table:
                    # 保存前一个表格
                    tables[current_table] = {
                        'headers': current_headers,
                        'rows': current_rows
                    }

                # 开始新表格
                table_name = f"表格_{len(tables) + 1}"
                if i > 0:
                    # 尝试使用前一行作为表格名称
                    prev_line = lines[i-1].strip()
                    if prev_line and len(prev_line) < 50:  # 避免使用过长的文本作为表名
                        table_name = prev_line

                current_table = table_name
                current_headers = self._parse_table_header(line)
                current_rows = []

                # 跳过表头行
                i += 1
                continue

            # 如果是数据行
            if current_table and line and not line.startswith('---'):
                # 尝试解析数据行
                row_data = self._parse_table_row(line, current_headers)
                if row_data:
                    current_rows.append(row_data)

            i += 1

        # 保存最后一个表格
        if current_table:
            tables[current_table] = {
                'headers': current_headers,
                'rows': current_rows
            }

        return tables

    def _parse_table_header(self, header_line: str) -> List[str]:
        """解析表格表头"""
        # 简单的表头解析，根据空格分割
        headers = []
        parts = re.split(r'\s{2,}', header_line)  # 两个或更多空格作为分隔符

        for part in parts:
            part = part.strip()
            if part:
                headers.append(part)

        return headers

    def _parse_table_row(self, row_line: str, headers: List[str]) -> Dict[str, str]:
        """解析表格数据行"""
        if not headers:
            return {}

        # 简单的行解析，根据空格分割
        parts = re.split(r'\s{2,}', row_line)  # 两个或更多空格作为分隔符

        row_data = {}
        for i, header in enumerate(headers):
            if i < len(parts):
                row_data[header] = parts[i].strip()
            else:
                row_data[header] = ""

        return row_data

    def _create_file_rule_from_table(self, table_name: str, table_data: Dict, is_expected: bool = False) -> FileRule:
        """从表格数据创建文件规则"""
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])

        # 提取列信息
        columns = []
        for header in headers:
            if header and header != '列名':
                columns.append(header)

        # 创建sheet规则
        sheet_rule = SheetRule(
            sheet_name="Sheet1",
            columns=columns,
            data_samples=rows[:3] if rows else []  # 取前3行作为数据示例
        )

        file_name = f"{table_name}.xlsx"
        if is_expected:
            file_name = f"expected_{file_name}"

        return FileRule(
            file_name=file_name,
            sheets=[sheet_rule]
        )

    def _extract_mapping_from_tables(self, tables: Dict[str, Dict]) -> Dict[str, str]:
        """从表格中提取映射规则"""
        mapping_rules = {}

        for table_name, table_data in tables.items():
            headers = table_data.get('headers', [])
            rows = table_data.get('rows', [])

            # 检查是否有"与报表项关系"列
            if '与报表项关系' in headers:
                relation_index = headers.index('与报表项关系')
                column_name_index = headers.index('列名') if '列名' in headers else 0

                for row in rows:
                    if len(row) > max(relation_index, column_name_index):
                        source = row[column_name_index]
                        target = row[relation_index]

                        if source and target and target != '是' and target != '否':
                            mapping_rules[source] = target

        return mapping_rules

    def _extract_mapping_rules(self, text: str) -> Dict[str, str]:
        """提取映射规则"""
        mapping_rules = {}

        # 查找映射规则部分
        mapping_patterns = [
            r'映射规则[：:]\s*\n((?:.+\n)+)',
            r'Mapping Rules[：:]\s*\n((?:.+\n)+)',
            r'数据映射[：:]\s*\n((?:.+\n)+)'
        ]

        for pattern in mapping_patterns:
            match = re.search(pattern, text, re.MULTILINE)
            if match:
                mapping_text = match.group(1)

                # 解析映射规则行
                lines = mapping_text.strip().split('\n')
                for line in lines:
                    if '->' in line or '→' in line:
                        parts = re.split(r'->|→', line)
                        if len(parts) == 2:
                            source = parts[0].strip()
                            target = parts[1].strip()
                            mapping_rules[source] = target

        return mapping_rules

    def _extract_calculation_rules(self, text: str) -> Dict[str, str]:
        """提取计算规则"""
        calculation_rules = {}

        # 查找计算规则部分
        calc_patterns = [
            r'计算规则[：:]\s*\n((?:.+\n)+)',
            r'Calculation Rules[：:]\s*\n((?:.+\n)+)',
            r'计算公式[：:]\s*\n((?:.+\n)+)'
        ]

        for pattern in calc_patterns:
            match = re.search(pattern, text, re.MULTILINE)
            if match:
                calc_text = match.group(1)

                # 解析计算规则行
                lines = calc_text.strip().split('\n')
                for line in lines:
                    if '=' in line:
                        parts = line.split('=', 1)
                        if len(parts) == 2:
                            column = parts[0].strip()
                            formula = parts[1].strip()
                            calculation_rules[column] = formula

        return calculation_rules


class WordRuleParser(RuleParser):
    """Word规则文件解析器"""

    def __init__(self):
        super().__init__()
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx未安装，无法使用Word解析功能")

    def parse(self, file_path: str) -> RuleSet:
        """解析Word规则文件"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Word文件不存在: {file_path}")

        # 提取Word文本
        text = self._extract_word_text(file_path)

        # 使用PDF解析器的文本解析逻辑
        pdf_parser = PDFRuleParser()
        return pdf_parser._parse_rules_from_text(text, file_path)

    def _extract_word_text(self, file_path: str) -> str:
        """从Word文件中提取文本"""
        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))

        return '\n'.join(text_parts)


class ExcelRuleParser(RuleParser):
    """Excel规则文件解析器"""

    def parse(self, file_path: str) -> RuleSet:
        """解析Excel规则文件"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Excel文件不存在: {file_path}")

        # 使用现有的Excel解析器
        sheet_data_list = self.excel_parser.parse_excel_file(file_path, active_sheet_only=True, max_data_rows=30)

        # 解析规则
        return self._parse_rules_from_excel(sheet_data_list, file_path)

    def _parse_rules_from_excel(self, sheet_data_list: List[Any], source_file: str) -> RuleSet:
        """从Excel数据中解析规则"""
        # 这里需要根据实际的Excel结构来解析规则
        # 假设规则文件有特定的结构

        expected_file_name = None
        source_file_rules = []
        mapping_rules = {}
        calculation_rules = {}

        for sheet_data in sheet_data_list:
            sheet_name = sheet_data.sheet_name

            if sheet_name.lower() in ['规则', 'rules', 'mapping', '映射']:
                # 解析规则表
                for region in sheet_data.regions:
                    for row in region.data:
                        # 解析规则行
                        self._parse_rule_row(row, region.head_data,
                                            mapping_rules, calculation_rules)

            elif sheet_name.lower() in ['预期输出', 'expected', 'output']:
                # 解析预期输出表
                expected_file_name = self._parse_expected_output(sheet_data)

        # 如果没有找到预期文件名，使用默认值
        if not expected_file_name:
            base_name = Path(source_file).stem
            expected_file_name = f"{base_name}_output.xlsx"

        # 创建规则集
        expected_file_rule = FileRule(
            file_name=expected_file_name,
            sheets=[]
        )

        return RuleSet(
            expected_file=expected_file_rule,
            source_files=source_file_rules,
            mapping_rules=mapping_rules,
            calculation_rules=calculation_rules
        )

    def _parse_rule_row(self, row: Dict[str, Any], headers: Dict[str, str],
                       mapping_rules: Dict[str, str], calculation_rules: Dict[str, str]):
        """解析规则行"""
        # 根据表头结构解析规则
        for header, column_letter in headers.items():
            value = row.get(column_letter)
            if value:
                header_lower = header.lower()

                if 'source' in header_lower or '来源' in header_lower:
                    # 数据来源
                    pass
                elif 'target' in header_lower or '目标' in header_lower:
                    # 目标列
                    pass
                elif 'formula' in header_lower or '公式' in header_lower or '计算' in header_lower:
                    # 计算公式
                    if '->' in str(value) or '→' in str(value):
                        # 映射规则
                        parts = re.split(r'->|→', str(value))
                        if len(parts) == 2:
                            source = parts[0].strip()
                            target = parts[1].strip()
                            mapping_rules[source] = target
                    elif '=' in str(value):
                        # 计算规则
                        parts = str(value).split('=', 1)
                        if len(parts) == 2:
                            column = parts[0].strip()
                            formula = parts[1].strip()
                            calculation_rules[column] = formula

    def _parse_expected_output(self, sheet_data: Any) -> Optional[str]:
        """解析预期输出信息"""
        for region in sheet_data.regions:
            for row in region.data:
                for header, column_letter in region.head_data.items():
                    value = row.get(column_letter)
                    if value and isinstance(value, str):
                        if 'file' in header.lower() or '文件' in header.lower():
                            return value.strip()
                        elif value.lower().endswith('.xlsx') or value.lower().endswith('.xls'):
                            return value.strip()
        return None


class RuleParserFactory:
    """规则解析器工厂"""

    @staticmethod
    def create_parser(file_path: str) -> RuleParser:
        """根据文件类型创建解析器"""
        file_ext = Path(file_path).suffix.lower()

        if file_ext == '.pdf':
            if not PDF_AVAILABLE:
                raise ImportError("PDF解析需要安装PyPDF2: pip install PyPDF2")
            return PDFRuleParser()
        elif file_ext in ['.docx', '.doc']:
            if not DOCX_AVAILABLE:
                raise ImportError("Word解析需要安装python-docx: pip install python-docx")
            return WordRuleParser()
        elif file_ext in ['.xlsx', '.xls']:
            return ExcelRuleParser()
        else:
            raise ValueError(f"不支持的文件类型: {file_ext}")


# 使用示例
if __name__ == "__main__":
    # 测试PDF解析
    try:
        parser = RuleParserFactory.create_parser("rules.pdf")
        rule_set = parser.parse("rules.pdf")
        print(f"解析成功: {rule_set}")
    except Exception as e:
        print(f"PDF解析测试失败: {e}")

    # 测试Word解析
    try:
        parser = RuleParserFactory.create_parser("rules.docx")
        rule_set = parser.parse("rules.docx")
        print(f"解析成功: {rule_set}")
    except Exception as e:
        print(f"Word解析测试失败: {e}")

    # 测试Excel解析
    try:
        parser = RuleParserFactory.create_parser("rules.xlsx")
        rule_set = parser.parse("rules.xlsx")
        print(f"解析成功: {rule_set}")
    except Exception as e:
        print(f"Excel解析测试失败: {e}")