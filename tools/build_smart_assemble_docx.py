from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "documents" / "智能组表概要设计_非固定数据源到固定报表.docx"

FONT = "Microsoft YaHei"
INK = "20252B"
BLUE = "2E5E9E"
HEADER_FILL = "EEF1F4"
GRID = "CCD2D9"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    set_keep_with_next(p)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        r1.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_numbered(doc, items, level=0):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.65 if level == 0 else 1.3)
        p.paragraph_format.first_line_indent = Cm(-0.65)
        p.add_run(f"{index}.  {item}")


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    set_row_cant_split(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Cm(widths_cm[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, 9, True, INK)
    for row in rows:
        added_row = table.add_row()
        set_row_cant_split(added_row)
        cells = added_row.cells
        for i, text in enumerate(row):
            cells[i].width = Cm(widths_cm[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(text))
            set_run_font(r, 9, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, 8, False, "777777")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35

    for name, size, color, before, after in [
        ("Heading 1", 18, INK, 0, 12),
        ("Heading 2", 15, INK, 14, 7),
        ("Heading 3", 12, BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Number", "List Number 2", "List Bullet"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)
    configure_styles(doc)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.add_run("智能组表概要设计")
    set_run_font(fp.runs[0], 8, False, "777777")
    fp.add_run("    ")
    add_page_number(fp)

    add_heading(doc, "6 智能组表", 1)
    add_body(doc, "智能组表用于实现“数据源不固定、输出报表格式固定”的自动化组装。用户每次可上传数量、文件名、Sheet 名、列顺序和字段命名均可能变化的 Excel 数据源，并指定一个固定输出模板。系统依据规则识别各数据表角色、人员主体和字段语义，将数据写入模板规定区域，最终输出结构、列顺序、表头和样式稳定的报表。")
    p = add_body(doc, "设计边界：本章仅描述智能组表能力，不扩展智算平台的通用训练与脚本计算模块。无法识别主键、源表存在严重结构损坏或规则本身冲突时，系统停止自动填充并提示人工处理。", "设计边界：")
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.right_indent = Cm(0.35)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)

    add_heading(doc, "6.1 系统功能说明", 2)
    add_table(doc, ["功能点", "功能描述", "产出"], [
        ("规则与模板约束", "选择全局或客户级组表规则，以用户上传模板的激活 Sheet 作为固定输出结构。", "规则上下文、模板结构"),
        ("非固定源解析", "支持多文件及全部可见 Sheet，识别表头、列字母、公式列、格式及少量脱敏样例。", "标准化源结构"),
        ("主体与字段匹配", "先确定人员集合及关联主键，再结合文件名、Sheet 名、列名和规则完成字段映射。", "结构化字段映射"),
        ("固定模板组装", "保留模板既有结构与样式，按需扩展数据行，将多源数据关联、合并后写入固定列。", "固定格式原版报表"),
        ("结果复核与复用", "记录任务映射；正确结果提升置信度，错误结果允许逐列修正并重新组表。", "纯值版、知识库、代码存档"),
    ], [3.4, 9.0, 4.2])
    add_body(doc, "适用范围：源文件允许按月份、地区或客户变化；输出模板在一次任务内固定，输出列、顺序、表头、样式和公式要求以模板为准。优先覆盖人事、薪资、社保、公积金等按人员主键汇总的二维明细表。", "适用范围：")

    add_heading(doc, "6.2 逻辑模型", 2)
    add_body(doc, "智能组表由组表规则、组表任务、字段映射知识三个核心对象组成，文件实体按租户隔离存储。")
    add_table(doc, ["对象", "关键属性", "说明"], [
        ("组表规则 assemble_rules", "名称、作用域、租户、附件、说明、上传人", "全局规则供所有租户选择；客户规则仅指定租户可见。"),
        ("组表任务 assemble_tasks", "租户、规则、结构签名、状态、代码路径、映射、输出、反馈", "记录一次从上传、分析、执行到复核的完整链路。"),
        ("字段映射 assemble_field_mappings", "租户、源列、目标列、模板签名、匹配类型、确认次数、状态", "沉淀同类模板下可复用的字段对应关系。"),
    ], [4.6, 7.0, 5.0])
    add_bullets(doc, [
        "一个任务对应一个租户、一个固定模板和一个或多个源文件，输出一个原版及可选纯值版。",
        "字段映射以 tenant_id + template_signature + source_column 唯一，避免跨客户、跨模板误复用。",
        "全局规则存放于 global_assets/assemble_rules/；客户规则、代码存档和任务结果按 tenants/{tenant_id}/ 隔离。",
    ])

    add_heading(doc, "6.3 系统功能权限", 2)
    add_table(doc, ["功能 / 角色", "普通业务用户", "规则管理员", "平台管理员"], [
        ("查看可用规则", "本租户 + 全局", "允许", "允许"),
        ("提交任务、查看进度、下载结果", "本租户", "本租户", "授权租户"),
        ("查看历史、反馈、修正映射重跑", "本人 / 本租户", "本租户", "授权租户"),
        ("维护组表规则", "不允许", "允许", "允许"),
        ("维护字段映射知识", "不允许", "允许", "允许"),
    ], [6.2, 3.5, 3.3, 3.6])
    add_body(doc, "权限控制要求：任务执行使用 tools.assemble；规则及知识库管理使用 tools.assemble.manage。任务、代码存档、字段映射、归档输入和输出文件均按 tenant_id 隔离。", "权限控制要求：")

    add_heading(doc, "6.4 接口定义", 2)
    add_table(doc, ["接口", "方法", "用途"], [
        ("/api/assemble/rules", "GET / POST", "查询可用规则；管理员上传全局或客户规则。"),
        ("/api/assemble/rules/{id}/upload", "POST", "替换指定规则附件。"),
        ("/api/assemble/submit", "POST", "提交源文件、固定模板、规则和强制重匹配标识。"),
        ("/api/assemble/tasks/{id}/stream", "GET", "SSE 返回解析、匹配、生成、执行和完成事件。"),
        ("/api/assemble/tasks/{id}/status", "GET", "查询任务状态、输出文件和错误信息。"),
        ("/api/assemble/download/{id}/{file}", "GET", "下载原版或纯值版结果。"),
        ("/api/assemble/history", "GET", "按授权租户查询任务历史。"),
        ("/api/assemble/tasks/{id}/feedback", "POST", "确认正确或标记有误，更新映射置信度。"),
        ("/api/assemble/tasks/{id}/rematch", "POST", "使用人工修正映射和归档输入重新组表。"),
        ("/api/assemble/mappings", "GET / PUT / DELETE", "管理员维护字段映射知识库。"),
    ], [7.6, 2.6, 6.4])

    add_heading(doc, "6.5 数据初始化", 2)
    add_numbered(doc, [
        "初始化权限点 tools.assemble 和 tools.assemble.manage，并分配至对应业务角色。",
        "初始化智能组表全局规则，至少包含组表模式、人员主体、关联主键、字段冲突和空值策略。",
        "创建 assemble_rules、assemble_tasks、assemble_field_mappings 三类数据表及唯一约束、索引。",
        "建立全局规则、租户规则、代码存档和任务结果目录，并设置容量、留存和清理策略。",
    ])

    add_heading(doc, "6.6 重点功能设计", 2)
    add_heading(doc, "6.6.1 整体业务流程", 3)
    add_numbered(doc, [
        "提交：选择规则，上传多个非固定源文件和一个固定输出模板；校验文件类型、租户权限并检测加密文件。",
        "解析：读取全部可见源 Sheet，仅以模板激活 Sheet 作为目标区域；进入 AI 上下文和日志的样例须先脱敏。",
        "签名：按源列集合、模板有效表头和规则内容生成稳定签名，文件名或月份变化不应导致无效失配。",
        "匹配：优先命中代码存档，其次采用已确认映射，再由 AI 处理剩余字段；歧义字段不得猜测。",
        "组装：先对齐人员主体，再按主键关联、纵向拼接或采用两阶段混合组装。",
        "输出：复制模板、填充数据、按需扩展样式行，生成原版和纯值版。",
        "复核：用户确认正确，或逐列修正映射后使用归档输入重新组表。",
    ])

    add_heading(doc, "6.6.2 非固定数据源解析与固定模板识别", 3)
    add_table(doc, ["项目", "非固定数据源", "固定输出模板"], [
        ("读取范围", "多个 Excel 文件、全部可见 Sheet", "单个模板文件、激活 Sheet"),
        ("提取内容", "文件名、Sheet 名、表头、列字母、公式、格式、前 3 行脱敏样例", "表头行、数据区域、列名、列字母、公式列和列格式"),
        ("变化容忍", "允许文件名、Sheet 名、列顺序和同义字段变化", "表头结构决定输出签名；Sheet 按月改名不影响识别"),
        ("异常策略", "无表头时使用解析兜底并记录日志", "只有表头无数据行时停止并明确提示"),
    ], [3.0, 6.8, 6.8])
    add_bullets(doc, [
        "源签名：对全部源文件的列名集合进行稳定排序后计算摘要。",
        "模板签名：基于激活 Sheet 的有效表头和关键列集合计算，不包含 Sheet 名。",
        "总签名：规则内容哈希 + 源签名 + 模板签名。命中且未强制重匹配时直接执行已验证代码。",
    ])

    add_heading(doc, "6.6.3 组表模式与字段匹配", 3)
    add_table(doc, ["组表模式", "识别特征", "处理方式"], [
        ("同构整合", "多个源表列结构高度相似，如不同地区同类明细。", "纵向拼接；同主键数值加总、文本取首值并去重。"),
        ("职责分工", "各源表字段差异大，分别提供花名册、考勤、社保等信息。", "确定人员主体后，按主键关联各表负责字段。"),
        ("混合模式", "部分表同构，另有入职、离职或专项补充表。", "先同构合并形成中间结果，再与职责表关联。"),
    ], [3.2, 6.6, 6.8])
    add_body(doc, "匹配原则：先确定人员主体和关联主键，再匹配字段。主键优先级为身份证/证件号 > 工号/员工编号 > 姓名+部门 > 姓名。无法确定主键或存在多人重名时应留空并告警，不得猜测填充。", "匹配原则：")

    add_heading(doc, "6.6.4 映射复用与置信度闭环", 3)
    add_table(doc, ["层级", "命中条件", "处理"], [
        ("代码存档", "规则、源结构和模板结构总签名一致", "跳过 AI，直接执行已验证代码。"),
        ("已确认语义映射", "同租户、同模板签名、源列一致，状态 active 且确认次数达标", "自动采用并记录 used_mapping_ids。"),
        ("同名列", "源列名与模板有效列名完全一致", "确定性白名单直接采用。"),
        ("AI 候选映射", "未被前述层级覆盖的字段", "成功执行后以 pending 状态入库。"),
        ("冲突或否定映射", "同源列指向多个目标，或用户标记错误", "置为 review_needed，后续不自动采用。"),
    ], [3.6, 7.2, 5.8])
    add_bullets(doc, [
        "语义映射累计 2 次人工确认后转为 active，避免一次错误结果污染后续任务。",
        "用户可逐列修正映射；修正后复用归档源文件和模板重跑，并将修正映射持久化到代码存档。",
    ])

    add_heading(doc, "6.6.5 执行与固定输出", 3)
    add_bullets(doc, [
        "执行代码在独立子进程和受控沙箱中运行，避免 Excel 计算阻塞 API 事件循环。",
        "复制固定模板后填充目标数据行；数据超过模板预留范围时复制最后数据行样式向下扩展。",
        "原版保留源数据 Sheet 和新生成公式；纯值版将新增公式计算为值并移除源 Sheet。",
        "结果按租户和任务落盘，下载时校验任务归属和文件名，禁止跨租户访问。",
    ])

    add_heading(doc, "6.6.6 异常、安全与验收", 3)
    add_table(doc, ["类别", "设计要求", "验收要点"], [
        ("数据安全", "仅脱敏样例进入 AI 上下文；真实数据仅在本地沙箱执行。", "姓名、证件、手机、邮箱、金额等日志样例不可明文泄露。"),
        ("执行异常", "AI 生成失败自动携带错误重试 1 次；执行失败记录 error 并支持重试。", "失败代码和未经验证映射不得写入存档或知识库。"),
        ("过程可见", "SSE 展示解析、匹配、生成、执行和完成状态。", "长任务日志持续可见，断线后可查询最终状态。"),
        ("固定输出", "输出列、顺序、表头、样式与模板一致，人员及金额按规则组装。", "同构整合、职责分工和混合场景均通过业务样例核对。"),
    ], [3.0, 7.5, 6.1])

    doc.core_properties.title = "智能组表概要设计 - 非固定数据源到固定报表"
    doc.core_properties.subject = "服务交付报表概要设计第6章"
    doc.core_properties.author = ""
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
