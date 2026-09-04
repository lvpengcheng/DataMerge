from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from build_smart_assemble_overview_docx import (
    FONT,
    MUTED,
    add_body,
    add_callout,
    add_heading,
    add_list,
    create_numbering,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "documents" / "智能组表概要设计_终端格式报表转换.docx"
OUTPUT = ROOT / "output" / "documents" / "智能组表概要设计_终端格式报表转换_优化版.docx"
DIAGRAM = ROOT / "tmp" / "smart_assemble_section_66" / "智能组表处理流程.png"


def _font(size, bold=False):
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def _centered_text(draw, box, text, font, fill="#344054"):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=8, align="center")
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(
        ((x1 + x2 - width) / 2, (y1 + y2 - height) / 2 - bbox[1]),
        text,
        font=font,
        fill=fill,
        spacing=8,
        align="center",
    )


def _box(draw, box, text, fill="#F2F4F7", outline="#667085", width=3, radius=16):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)
    _centered_text(draw, box, text, _font(30), fill="#344054")


def _arrow(draw, start, end, color="#667085", width=4, label=None):
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        head = [(x2, y2), (x2 - 15 * direction, y2 - 10), (x2 - 15 * direction, y2 + 10)]
    else:
        direction = 1 if y2 > y1 else -1
        head = [(x2, y2), (x2 - 10, y2 - 15 * direction), (x2 + 10, y2 - 15 * direction)]
    draw.polygon(head, fill=color)
    if label:
        lx = (x1 + x2) / 2 + (8 if x1 == x2 else 0)
        ly = (y1 + y2) / 2 - (30 if y1 == y2 else 0)
        draw.text((lx, ly), label, font=_font(25), fill="#475467", anchor="mm")


def build_diagram():
    DIAGRAM.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (2100, 1420), "white")
    draw = ImageDraw.Draw(image)
    draw.text((1050, 55), "智能组表处理流程", font=_font(42, True), fill="#24292F", anchor="mm")

    w, h = 470, 120
    xs = [130, 815, 1500]
    ys = [145, 445, 755, 1065]
    boxes = {}

    labels = [
        ("upload", 0, 0, "上传源文件和\n目标模板", "#EAF2FB", "#4B82BD"),
        ("rules", 1, 0, "加载全局规则和\n个性化规则", "#EAF2FB", "#4B82BD"),
        ("parse", 2, 0, "解析文件结构、字段\n及场景特征", "#EAF2FB", "#4B82BD"),
        ("script", 0, 1, "已有脚本可复用？", "#FFF2CC", "#D6A400"),
        ("run", 1, 1, "执行已归档脚本", "#E6FAF7", "#36AFA1"),
        ("output_fast", 2, 1, "生成终端格式报表", "#E6F6EA", "#4BAE63"),
        ("mapping", 0, 2, "查询系统已有\n字段映射", "#F2F4F7", "#667085"),
        ("ai", 1, 2, "AI 分析未匹配字段", "#EEE7FF", "#805AD5"),
        ("save", 2, 2, "确认结果并写入\n字段映射表", "#E6FAF7", "#36AFA1"),
        ("assemble", 0, 3, "按完整映射关系\n组装报表", "#F2F4F7", "#667085"),
        ("archive", 1, 3, "生成并归档\nPython 脚本", "#E6FAF7", "#36AFA1"),
        ("output", 2, 3, "生成终端格式报表", "#E6F6EA", "#4BAE63"),
    ]
    for key, col, row, text, fill, outline in labels:
        box = (xs[col], ys[row], xs[col] + w, ys[row] + h)
        boxes[key] = box
        _box(draw, box, text, fill=fill, outline=outline)

    def right(key):
        x1, y1, x2, y2 = boxes[key]
        return (x2, (y1 + y2) // 2)

    def left(key):
        x1, y1, x2, y2 = boxes[key]
        return (x1, (y1 + y2) // 2)

    def top(key):
        x1, y1, x2, y2 = boxes[key]
        return ((x1 + x2) // 2, y1)

    def bottom(key):
        x1, y1, x2, y2 = boxes[key]
        return ((x1 + x2) // 2, y2)

    _arrow(draw, right("upload"), left("rules"))
    _arrow(draw, right("rules"), left("parse"))
    # Return to the next row without a long horizontal chain.
    px, py = bottom("parse")
    sx, sy = top("script")
    draw.line([(px, py), (px, sy - 45), (sx, sy - 45), (sx, sy)], fill="#667085", width=4)
    draw.polygon([(sx, sy), (sx - 10, sy - 15), (sx + 10, sy - 15)], fill="#667085")
    _arrow(draw, right("script"), left("run"), label="是")
    _arrow(draw, right("run"), left("output_fast"))
    _arrow(draw, bottom("script"), top("mapping"), label="否")
    _arrow(draw, right("mapping"), left("ai"), label="未匹配")
    _arrow(draw, right("ai"), left("save"))
    sx, sy = bottom("save")
    tx, ty = top("assemble")
    draw.line([(sx, sy), (sx, ty - 45), (tx, ty - 45), (tx, ty)], fill="#667085", width=4)
    draw.polygon([(tx, ty), (tx - 10, ty - 15), (tx + 10, ty - 15)], fill="#667085")
    _arrow(draw, right("assemble"), left("archive"))
    _arrow(draw, right("archive"), left("output"))

    draw.text(
        (1050, 1340),
        "已有映射优先复用；AI 仅处理未匹配字段，确认后沉淀为映射与脚本。",
        font=_font(27),
        fill="#667085",
        anchor="mm",
    )
    image.save(DIAGRAM, dpi=(220, 220))


def remove_after_heading(document, heading_text):
    body = document._element.body
    target = None
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == heading_text:
            target = paragraph._p
            break
    if target is None:
        raise ValueError(f"未找到章节标题：{heading_text}")
    found = False
    for child in list(body):
        if child is target:
            found = True
            continue
        if found and child.tag != qn("w:sectPr"):
            body.remove(child)


def isolate_numbering(document, abstract_id, signature):
    """Prevent Word from merging visually adjacent lists across subsections."""
    numbering = document.part.numbering_part.element
    for abstract in numbering.findall(qn("w:abstractNum")):
        if abstract.get(qn("w:abstractNumId")) == str(abstract_id):
            nsid = OxmlElement("w:nsid")
            nsid.set(qn("w:val"), signature)
            abstract.insert(0, nsid)
            tmpl = OxmlElement("w:tmpl")
            tmpl.set(qn("w:val"), signature[::-1])
            abstract.insert(1, tmpl)
            return
    raise ValueError(f"未找到编号定义：{abstract_id}")


def add_process_figure(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run()
    run.add_picture(str(DIAGRAM), width=Cm(15.6))
    drawing = run._r.find(qn("w:drawing"))
    if drawing is not None:
        doc_pr = drawing.find(".//" + qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", "智能组表处理流程图")


def build():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    build_diagram()
    document = Document(SOURCE)
    remove_after_heading(document, "6.6 重点功能设计")

    numbering, decimal_abs, bullet_abs, next_id = create_numbering(document)
    isolate_numbering(document, decimal_abs, "66A10001")
    numbering_state = [numbering, decimal_abs, bullet_abs, next_id]

    add_body(
        document,
        "前期验证表明，不同租户的计算口径、组表条件和数据源结构难以统一，AI 无法在缺少明确规则的情况下自动推导复杂业务计算并保证结果完全准确。因此，本阶段不建设面向任意数据源的通用报表生成能力，重点建设边界明确的终端格式报表转换能力。",
    )

    add_heading(document, "6.6.1 整体业务流程", 3)
    add_body(
        document,
        "智能组表以已完成业务计算的源文件为输入，以固定终端报表模板为输出，按照“规则加载—匹配复用—AI 补充—映射落表—报表组装—脚本归档”的流程执行。",
    )
    add_process_figure(document)
    add_list(
        document,
        [
            "任务发起：用户选择试点场景，上传一个或多个源文件及固定终端格式模板。",
            "规则加载：系统加载通用字段识别、格式处理等全局规则，并叠加租户或功能对应的个性化规则。",
            "文件解析：识别源文件和模板的 Sheet、表头、字段、数据区域及场景特征。",
            "复用判断：优先匹配已归档的 Python 脚本；场景特征和规则版本一致时直接执行。",
            "字段匹配：无可复用脚本时，先查询系统已有字段映射；仅对未匹配字段调用 AI 进行同名或同语义分析。",
            "映射沉淀：AI 分析结果经规则校验或业务确认后写入字段映射表，形成完整映射关系。",
            "报表生成：系统按映射关系完成字段搬运、固定规则处理和模板组装，生成终端格式报表，同时生成并归档可复用的 Python 脚本。",
        ],
        numbering_state,
        ordered=True,
    )

    add_heading(document, "6.6.2 智能组表处理", 3)
    add_body(
        document,
        "智能组表不重新计算薪资、社保、服务费等业务结果，主要完成已有数据的识别、匹配、搬运和必要的固定处理。核心功能如下：",
    )
    # Use a distinct numbering definition for each functional subsection so
    # Microsoft Word reliably restarts the visible sequence at 1.
    numbering2, decimal_abs2, bullet_abs2, next_id2 = create_numbering(document)
    isolate_numbering(document, decimal_abs2, "66A20002")
    numbering_state2 = [numbering2, decimal_abs2, bullet_abs2, next_id2]
    add_list(
        document,
        [
            "规则管理：全局规则用于定义通用识别和处理要求；个性化规则用于补充租户或功能特有的字段来源、筛选、合并、排序及简单转换要求。",
            "映射复用：同一场景优先使用字段映射表中的已确认关系，减少重复分析并保证多次执行结果一致。",
            "AI 辅助匹配：当系统中没有匹配关系时，AI 根据字段名称、字段语义、样例数据和规则说明给出候选映射；无法可靠判断时提示人工确认，不自动猜测。",
            "报表组装：按照完整映射关系将源数据写入目标 Sheet 和字段位置，保持模板表头、字段顺序及基本格式，并执行已明确的多文件合并、筛选和排序规则。",
            "成果固化：将验证通过的映射关系和转换逻辑生成 Python 脚本并按场景、模板及规则版本归档，供后续任务直接复用。",
        ],
        numbering_state2,
        ordered=False,
    )
    add_callout(
        document,
        "功能边界：",
        "本功能仅处理已明确规则下的数据转换，不由 AI 自主推导复杂业务计算口径；输入结构、目标模板或个性化规则发生实质变化时，应重新匹配、验证并发布脚本版本。",
    )

    add_heading(document, "6.6.3 客户端任务管理", 3)
    numbering3, decimal_abs3, bullet_abs3, next_id3 = create_numbering(document)
    isolate_numbering(document, decimal_abs3, "66A30003")
    numbering_state3 = [numbering3, decimal_abs3, bullet_abs3, next_id3]
    add_list(
        document,
        [
            "任务创建：用户选择已授权的试点场景、源文件和目标模板后发起组表任务。",
            "执行管理：展示解析、匹配、组装和生成状态；对缺少文件、字段无法匹配、规则冲突等情况给出明确提示并停止错误填充。",
            "结果管理：任务完成后提供终端格式报表保存或下载，并记录使用的映射版本、脚本版本和处理结果。",
            "版本更新：输入结构、模板或规则调整后，由场景管理员重新验证并发布新版本；同一版本资料重复执行应得到一致结果。",
        ],
        numbering_state3,
        ordered=False,
    )

    document.core_properties.title = "智能组表概要设计 - 终端格式报表转换（优化版）"
    document.core_properties.subject = "服务交付报表效能优化概要设计补充章节"
    document.core_properties.author = ""
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
