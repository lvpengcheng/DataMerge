from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "documents" / "智能组表概要设计_终端格式报表转换.docx"

# standard_business_brief preset with named overrides matching the supplied A4
# Chinese overview-design reference: A4 page, Microsoft YaHei, black headings.
FONT = "Microsoft YaHei"
INK = "24292F"
BLUE = "2156D9"
MUTED = "667085"
GRID = "D0D5DD"
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F7F8FA"
CONTENT_DXA = 9027
TABLE_INDENT_DXA = 120
CELL_START_DXA = 120


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    for attr in ("eastAsia", "ascii", "hAnsi"):
        fonts.set(qn(f"w:{attr}"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size, bold=False, color=INK):
    style.font.name = FONT
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def add_left_border(paragraph, color="C7CDD4", size="10", space="8"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)
    pbdr.append(left)


def shade_paragraph(paragraph, fill=CALLOUT_FILL):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def create_numbering(document):
    numbering = document.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    decimal_abs = max(existing_abs or [0]) + 1
    bullet_abs = decimal_abs + 1

    def abstract_num(abs_id, fmt, text, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "120")
        spacing.set(qn("w:line"), "264")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.extend([tabs, ind, spacing])
        lvl.extend([start, num_fmt, lvl_text, suff, ppr])
        if font:
            rpr = OxmlElement("w:rPr")
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:ascii"), font)
            rfonts.set(qn("w:hAnsi"), font)
            rpr.append(rfonts)
            lvl.append(rpr)
        abstract.append(lvl)
        numbering.append(abstract)

    abstract_num(decimal_abs, "decimal", "%1.")
    abstract_num(bullet_abs, "bullet", "•", FONT)
    next_num = max(existing_num or [0]) + 1
    return numbering, decimal_abs, bullet_abs, next_num


def new_num_id(numbering, abstract_id, next_id, restart=False):
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    if restart:
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        override.append(start_override)
        num.append(override)
    numbering.append(num)
    return next_id


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, numid])
    ppr.append(numpr)


def add_list(document, items, numbering_state, ordered=False):
    numbering, decimal_abs, bullet_abs, next_id = numbering_state
    num_id = new_num_id(numbering, decimal_abs if ordered else bullet_abs, next_id, restart=ordered)
    for item in items:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.1
        apply_num(p, num_id)
        p.add_run(item)
    numbering_state[3] += 1


def set_cell_margins(cell, top=80, bottom=80, start=CELL_START_DXA, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.find(qn("w:tcMar"))
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tcmar.append(node)


def set_table_geometry(table, widths):
    table.autofit = False
    tblpr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblpr.append(layout)
    tblw = tblpr.find(qn("w:tblW"))
    tblw.set(qn("w:type"), "dxa")
    tblw.set(qn("w:w"), str(sum(widths)))
    tblind = OxmlElement("w:tblInd")
    tblind.set(qn("w:type"), "dxa")
    tblind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tblpr.append(tblind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        trpr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trpr.append(cant)
        for idx, cell in enumerate(row.cells):
            tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tcw.set(qn("w:type"), "dxa")
            tcw.set(qn("w:w"), str(widths[idx]))


def set_table_borders(table):
    tblpr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "5")
        node.set(qn("w:color"), GRID)
        borders.append(node)
    tblpr.append(borders)


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0]
    header._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), HEADER_FILL)
        cell._tc.get_or_add_tcPr().append(shd)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        set_run_font(p.add_run(text), 9.5, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(str(text)), 9.5, False)
    set_table_geometry(table, widths)
    set_table_borders(table)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(document, text, level):
    p = document.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_body(document, text, lead=None):
    p = document.add_paragraph()
    if lead and text.startswith(lead):
        set_run_font(p.add_run(lead), 11, True)
        p.add_run(text[len(lead):])
    else:
        p.add_run(text)
    return p


def add_callout(document, label, text):
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.right_indent = Cm(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    add_left_border(p)
    shade_paragraph(p)
    set_run_font(p.add_run(label), 10.5, True)
    set_run_font(p.add_run(text), 10.5, False, MUTED)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, 8.5, False, MUTED)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    set_style_font(normal, 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    set_style_font(doc.styles["Heading 1"], 18, True)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(0)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(14)
    set_style_font(doc.styles["Heading 2"], 14, True)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(14)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(7)
    set_style_font(doc.styles["Heading 3"], 12, False)
    doc.styles["Heading 3"].paragraph_format.space_before = Pt(10)
    doc.styles["Heading 3"].paragraph_format.space_after = Pt(5)

    footer = section.footer.paragraphs[0]
    footer.add_run("智能组表概要设计    ")
    for r in footer.runs:
        set_run_font(r, 8.5, False, MUTED)
    add_page_number(footer)

    numbering, decimal_abs, bullet_abs, next_id = create_numbering(doc)
    numbering_state = [numbering, decimal_abs, bullet_abs, next_id]

    add_heading(doc, "6 智能组表", 1)
    add_body(doc, "智能组表面向已选定的试点租户和明确的报表转换场景，利用 AI 对输入源文件、目标输出模板及租户/功能计算逻辑说明进行分析，生成可复用的 Python AI Skill 和客户端可执行小工具，将源文件中的同名或同语义字段搬运至指定终端格式报表。")
    add_callout(doc, "设计定位：", "本功能不是一套可自动理解任意业务计算逻辑的通用报表平台，而是针对边界明确、材料齐备的试点场景，将一次性分析成果固化为确定性的转换工具。")

    add_heading(doc, "6.1 系统功能说明", 2)
    add_table(doc, ["应用系统", "功能点", "描述"], [
        ("服务交付报表平台", "统一交付报表可行性分析", "利用 AI 统一分析现有报表格式和数据源，识别自动化边界，形成不采用通用全自动生成方案的结论。"),
        ("智能组表", "终端格式报表转换", "针对试点租户，以已确认的输入源、目标模板和规则说明为依据，生成并使用客户端小工具完成字段搬运和格式转换。"),
    ], [1900, 2200, 4927])

    add_heading(doc, "6.2 逻辑模型", 2)
    add_body(doc, "智能组表围绕“场景资料包—转换能力—执行结果”组织。每个试点场景对应一套经过业务确认的资料和一项可复用的转换能力。")
    add_list(doc, [
        "场景资料包：试点租户范围、输入源文件、目标输出模板、租户/功能计算逻辑说明及业务确认样例。",
        "转换能力：AI 分析资料后生成的 Python AI Skill、转换脚本及客户端可执行小工具。",
        "执行结果：用户选择源文件后，由小工具按已固化规则生成终端格式报表，并提供处理结果及异常提示。",
    ], numbering_state)
    add_callout(doc, "核心关系：", "一个试点场景对应一套资料包和一个可执行转换工具；场景的输入结构、输出模板或规则发生实质变化时，应重新分析、验证并发布对应版本。")

    add_heading(doc, "6.3 系统功能权限", 2)
    add_table(doc, ["功能 / 角色", "服务交付人员", "场景管理员", "平台管理员"], [
        ("使用智能组表小工具", "授权试点范围", "允许", "允许"),
        ("维护试点场景资料", "不允许", "允许", "允许"),
        ("发布或更新转换工具", "不允许", "参与验收", "允许"),
    ], [3300, 1900, 1900, 1927])

    add_heading(doc, "6.4 接口定义", 2)
    add_body(doc, "无。当前阶段以客户端可执行小工具方式提供能力，输入和输出均为本地文件，不新增对外业务接口。后续如纳入统一平台运行，再结合实际使用情况补充接口设计。")

    add_heading(doc, "6.5 数据初始化", 2)
    add_body(doc, "试点场景上线前，由项目组和业务人员共同完成以下资料初始化：")
    add_list(doc, [
        "确定试点租户、适用功能和可使用人员范围。",
        "提供能够覆盖典型场景的输入源文件，并说明各文件用途。",
        "提供固定的终端格式输出模板及业务确认的正确结果样例。",
        "按租户或功能维护《每租户或者功能计算逻辑.md》，说明字段来源、筛选条件、合并规则及必要的简单转换要求。",
        "完成数据脱敏、样例完整性检查和业务验收基线确认。",
    ], numbering_state, ordered=True)

    add_heading(doc, "6.6 重点功能设计", 2)
    add_heading(doc, "6.6.1 AI 统一生成交付报表的可行性结论", 3)
    add_body(doc, "项目首先尝试利用 AI 同时分析交付报表格式和生成报表所需的数据源，以实现统一交付报表的自动生成。分析结果表明，现阶段无法仅依靠 AI 实现面向所有租户、所有报表场景的全自动生成并保证结果 100% 准确。主要原因如下：")
    add_list(doc, [
        "不同租户、不同功能的计算口径和业务规则不一致，部分规则依赖业务经验或线下约定。",
        "组表条件不一致，人员范围、筛选条件、关联主键、优先级和例外处理均可能不同。",
        "数据源的组成、数量和结构无法统一固定，同一报表在不同月份也可能发生变化。",
        "复杂计算逻辑需要明确且可验证的业务规则，AI 只能辅助理解，不能替代业务确认，也不能通过推测保证计算结果完全正确。",
    ], numbering_state)
    add_callout(doc, "结论：", "不将“AI 自动理解任意数据源并直接生成统一交付报表”作为本阶段建设目标。统一交付报表仍应依赖标准数据口径和明确的业务计算规则。")

    add_heading(doc, "6.6.2 智能组表的功能定位", 3)
    add_body(doc, "智能组表承接的是第二类需求：将交付数据转换为指定终端格式的报表。该场景不重新计算薪资、社保或服务费等业务结果，只对已经存在的数据进行识别、匹配、搬运和必要的简单整理。")
    add_list(doc, [
        "相同字段直接映射，例如“姓名”搬运至“姓名”。",
        "相同语义字段进行映射，例如“员工姓名”搬运至“姓名”、“实发合计”搬运至“实发工资”。",
        "按已经明确的固定规则完成多文件合并、字段筛选、顺序调整、名称转换和格式保持。",
        "将验证通过的转换逻辑固化为 Python 脚本和客户端工具，后续运行不再由 AI 临场推测。",
    ], numbering_state)

    add_heading(doc, "6.6.3 整体业务流程", 3)
    add_list(doc, [
        "确定试点场景：明确试点租户、适用报表、使用人员及范围边界。",
        "准备场景资料：收集输入源文件、目标输出模板、正确结果样例和《每租户或者功能计算逻辑.md》。",
        "AI 分析与生成：AI 分析文件结构、字段语义和规则说明，生成场景专用的 Python AI Skill 及转换脚本。",
        "工具封装：将验证通过的脚本封装为客户端可执行小工具，提供文件选择、执行、结果保存和异常提示能力。",
        "业务验证：使用典型样例和边界样例与正确结果逐项比对，经业务确认后发布。",
        "日常使用：服务交付人员在授权范围内选择源文件，执行小工具并生成终端格式报表。",
        "变更维护：输入结构、输出模板或规则发生变化时，更新资料、重新生成并完成回归验证后发布新版本。",
    ], numbering_state, ordered=True)

    add_heading(doc, "6.6.4 输入与输出", 3)
    add_table(doc, ["类别", "内容", "说明"], [
        ("输入", "试点租户的一个或多个源文件", "来源和用途在场景资料中明确，数据为已经完成业务计算的结果。"),
        ("输入", "固定终端格式模板", "明确目标 Sheet、表头、字段顺序、格式和输出要求。"),
        ("输入", "《每租户或者功能计算逻辑.md》", "记录字段对应、筛选、合并及简单转换规则，作为生成和验收依据。"),
        ("输出", "终端格式报表", "保持目标模板结构，将源数据按已确认规则写入指定位置。"),
        ("输出", "处理结果及异常提示", "提示未识别字段、缺少文件、主键无法匹配等情况，不静默猜测。"),
    ], [1300, 3200, 4527])

    add_heading(doc, "6.6.5 功能范围与边界", 3)
    add_table(doc, ["范围内", "范围外"], [
        ("同名或同语义字段匹配与搬运", "由 AI 自主推导租户复杂业务计算口径"),
        ("已明确规则下的多文件汇总、筛选和排序", "在规则缺失时猜测人员范围、关联关系或计算结果"),
        ("保持固定终端模板的字段顺序和基本格式", "面向所有租户、所有数据源的通用一次配置"),
        ("按试点场景生成并发布确定性客户端工具", "替代业务人员对正确结果的确认和验收"),
    ], [4513, 4514])

    add_heading(doc, "6.6.6 验收原则", 3)
    add_list(doc, [
        "覆盖选定试点租户的典型月份、多人、多文件及字段缺失等场景。",
        "生成结果与业务确认的正确结果样例逐字段、逐记录比对一致。",
        "同一版本资料重复执行应得到一致结果，不能因 AI 随机判断产生差异。",
        "缺少必要输入、字段无法匹配或规则冲突时应明确提示并停止错误填充。",
        "输入结构、输出模板或规则变更后，必须重新验证并发布新版本。",
    ], numbering_state)

    doc.core_properties.title = "智能组表概要设计 - 终端格式报表转换"
    doc.core_properties.subject = "服务交付报表效能优化概要设计补充章节"
    doc.core_properties.author = ""
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
